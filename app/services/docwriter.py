"""DocWriter service — DOCX and XLSX generation for proposals and compliance matrices."""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

logger = logging.getLogger(__name__)

# Brand colours
BRAND_DARK = RGBColor(0x1B, 0x3A, 0x5C)
BRAND_ACCENT = RGBColor(0x2E, 0x86, 0xC1)
HEADER_BG = "1B3A5C"
HEADER_FG = "FFFFFF"
SUBTLE_BG = "EBF5FB"


class DocWriterService:
    """Generates DOCX proposals and XLSX spreadsheets."""

    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # DOCX helpers
    # ------------------------------------------------------------------

    def _init_doc(self, title: str) -> Document:
        doc = Document()

        # Default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Heading styles
        for level in range(1, 4):
            heading_style = doc.styles[f"Heading {level}"]
            heading_style.font.color.rgb = BRAND_DARK
            heading_style.font.bold = True
            if level == 1:
                heading_style.font.size = Pt(18)
            elif level == 2:
                heading_style.font.size = Pt(14)
            else:
                heading_style.font.size = Pt(12)

        return doc

    def _add_cover_page(self, doc: Document, title: str, metadata: dict) -> None:
        """Add a professional cover page."""
        for _ in range(6):
            doc.add_paragraph("")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(28)
        run.font.color.rgb = BRAND_DARK
        run.bold = True

        doc.add_paragraph("")

        if metadata.get("client"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Prepared for: {metadata['client']}")
            run.font.size = Pt(14)
            run.font.color.rgb = BRAND_ACCENT

        if metadata.get("company"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Prepared by: {metadata['company']}")
            run.font.size = Pt(14)
            run.font.color.rgb = BRAND_ACCENT

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(datetime.now().strftime("%B %Y"))
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        if metadata.get("rfp_number"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"RFP Reference: {metadata['rfp_number']}")
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_page_break()

    def _add_toc_placeholder(self, doc: Document) -> None:
        """Add a Table of Contents placeholder heading."""
        doc.add_heading("Table of Contents", level=1)
        p = doc.add_paragraph("[Table of Contents — update field after opening in Word]")
        p.style = doc.styles["Normal"]
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        doc.add_page_break()

    def _format_table(self, table) -> None:
        """Apply consistent table formatting."""
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header row
        if table.rows:
            for cell in table.rows[0].cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
                shading = cell._element.get_or_add_tcPr()
                from docx.oxml.ns import qn
                shading_elem = shading.makeelement(qn("w:shd"), {
                    qn("w:fill"): HEADER_BG,
                    qn("w:val"): "clear",
                })
                shading.append(shading_elem)

    # ------------------------------------------------------------------
    # Technical Proposal
    # ------------------------------------------------------------------

    # Sections that appear in the opening pages before the TOC
    FRONT_MATTER_SECTIONS = {"Company Profile", "Past Successful Projects"}

    def _add_section_content(self, doc: Document, content: str) -> None:
        """Add section content to the document, handling markdown-style headings."""
        for para_text in content.split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            if para_text.startswith("### "):
                doc.add_heading(para_text.lstrip("# "), level=3)
            elif para_text.startswith("## "):
                doc.add_heading(para_text.lstrip("# "), level=2)
            else:
                doc.add_paragraph(para_text)

    def create_technical_proposal(
        self, title: str, sections: list[dict], metadata: dict
    ) -> str:
        """Create a full technical proposal DOCX.

        The document structure matches standard government tender submissions:
          1. Cover page
          2. Company Profile (front matter — before TOC)
          3. Past Successful Projects (front matter — before TOC)
          4. Table of Contents
          5. Executive Summary, Technical Approach, etc.

        Returns path to file.
        """
        doc = self._init_doc(title)
        self._add_cover_page(doc, title, metadata)

        # Split sections into front matter and body
        front_matter = []
        body_sections = []
        for section in sections:
            if section["title"] in self.FRONT_MATTER_SECTIONS:
                front_matter.append(section)
            else:
                body_sections.append(section)

        # Front matter: Company Profile and Past Successful Projects
        # These appear right after the cover page, before the TOC
        for section in front_matter:
            doc.add_heading(section["title"], level=1)
            self._add_section_content(doc, section["content"])
            doc.add_page_break()

        # Table of Contents (after front matter, before technical body)
        self._add_toc_placeholder(doc)

        # Body sections: Executive Summary, Technical Approach, etc.
        for section in body_sections:
            doc.add_heading(section["title"], level=1)
            self._add_section_content(doc, section["content"])

        # Footer with page numbers
        section = doc.sections[-1]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{metadata.get('company', 'TenderAI')} — Confidential")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        filename = f"technical_proposal_{metadata.get('rfp_id', 'draft')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = self.output_dir / filename
        doc.save(str(output_path))
        logger.info("Created technical proposal: %s", output_path)
        return str(output_path)

    # ------------------------------------------------------------------
    # Compliance Matrix
    # ------------------------------------------------------------------

    def create_compliance_matrix(
        self, requirements: list[dict], responses: list[dict]
    ) -> str:
        """Create a compliance matrix DOCX. Returns path to file."""
        doc = self._init_doc("Compliance Matrix")
        doc.add_heading("Compliance Matrix", level=1)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["#", "Requirement", "Compliance Status", "Response"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for idx, (req, resp) in enumerate(zip(requirements, responses), 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = req.get("requirement", str(req))
            row.cells[2].text = resp.get("status", "Compliant")
            row.cells[3].text = resp.get("narrative", "")

        self._format_table(table)

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(0.5)
            row.cells[1].width = Inches(2.5)
            row.cells[2].width = Inches(1.5)
            row.cells[3].width = Inches(3.0)

        filename = f"compliance_matrix_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = self.output_dir / filename
        doc.save(str(output_path))
        logger.info("Created compliance matrix: %s", output_path)
        return str(output_path)

    # ------------------------------------------------------------------
    # Financial Proposal
    # ------------------------------------------------------------------

    def create_financial_proposal(
        self, bom_items: list[dict], metadata: dict
    ) -> str:
        """Create a financial proposal DOCX with pricing tables. Returns path to file."""
        doc = self._init_doc("Financial Proposal")
        self._add_cover_page(doc, f"Financial Proposal — {metadata.get('title', '')}", metadata)

        # Pricing summary
        doc.add_heading("Pricing Summary", level=1)

        # Group items by category
        categories: dict[str, list[dict]] = {}
        for item in bom_items:
            cat = item.get("category", "General")
            categories.setdefault(cat, []).append(item)

        grand_total = 0.0
        currency = metadata.get("currency", "OMR")

        for cat_name, items in categories.items():
            doc.add_heading(cat_name, level=2)
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            headers = ["Item", "Qty", "Unit Cost", "Margin", "Total"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h

            cat_total = 0.0
            for item in items:
                row = table.add_row()
                row.cells[0].text = item.get("item_name", "")
                row.cells[1].text = f"{item.get('quantity', 1):.0f}"
                row.cells[2].text = f"{currency} {item.get('unit_cost', 0):,.2f}"
                row.cells[3].text = f"{item.get('margin_pct', 15):.1f}%"
                total = item.get("total_cost", 0) or 0
                row.cells[4].text = f"{currency} {total:,.2f}"
                cat_total += total

            # Category subtotal row
            row = table.add_row()
            row.cells[0].text = f"{cat_name} Subtotal"
            row.cells[4].text = f"{currency} {cat_total:,.2f}"
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

            self._format_table(table)
            grand_total += cat_total

        # Grand total
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run(f"Grand Total: {currency} {grand_total:,.2f}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = BRAND_DARK

        # Terms
        doc.add_heading("Terms and Conditions", level=1)
        terms = [
            f"All prices are quoted in {currency} and are exclusive of VAT unless stated otherwise.",
            "This quotation is valid for 90 days from the date of submission.",
            "Payment terms: 30% upon contract signing, 50% upon delivery, 20% upon acceptance.",
            "Warranty: As specified per line item. Standard warranty is 12 months from acceptance.",
            "Delivery timelines are subject to confirmation upon contract award.",
        ]
        for term in terms:
            doc.add_paragraph(term, style="List Bullet")

        filename = f"financial_proposal_{metadata.get('rfp_id', 'draft')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = self.output_dir / filename
        doc.save(str(output_path))
        logger.info("Created financial proposal: %s", output_path)
        return str(output_path)

    # ------------------------------------------------------------------
    # BOM Spreadsheet
    # ------------------------------------------------------------------

    def create_bom_spreadsheet(self, bom_items: list[dict], metadata: dict) -> str:
        """Create a BOM XLSX spreadsheet. Returns path to file."""
        wb = Workbook()
        ws = wb.active
        ws.title = "Bill of Materials"

        # Styles
        header_font = Font(bold=True, color=HEADER_FG, size=11, name="Calibri")
        header_fill = PatternFill(start_color=HEADER_BG, end_color=HEADER_BG, fill_type="solid")
        currency = metadata.get("currency", "OMR")
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        # Title row
        ws.merge_cells("A1:I1")
        ws["A1"] = f"Bill of Materials — {metadata.get('title', 'Proposal')}"
        ws["A1"].font = Font(bold=True, size=14, color=HEADER_BG, name="Calibri")
        ws["A1"].alignment = Alignment(horizontal="center")

        # Headers
        headers = [
            "Category", "Item", "Description", "Manufacturer",
            "Part Number", "Qty", f"Unit Cost ({currency})",
            f"Margin %", f"Total ({currency})",
        ]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="center", wrap_text=True)

        # Data rows
        row_num = 4
        for item in bom_items:
            values = [
                item.get("category", ""),
                item.get("item_name", ""),
                item.get("description", ""),
                item.get("manufacturer", ""),
                item.get("part_number", ""),
                item.get("quantity", 1),
                item.get("unit_cost", 0),
                item.get("margin_pct", 15),
                item.get("total_cost", 0) or 0,
            ]
            for col, value in enumerate(values, 1):
                cell = ws.cell(row=row_num, column=col, value=value)
                cell.border = thin_border
                if col in (6, 7, 8, 9):
                    cell.number_format = "#,##0.00"
                    cell.alignment = Alignment(horizontal="right")
            row_num += 1

        # Grand total row
        total = sum((item.get("total_cost", 0) or 0) for item in bom_items)
        ws.cell(row=row_num + 1, column=8, value="Grand Total:").font = Font(bold=True, name="Calibri")
        total_cell = ws.cell(row=row_num + 1, column=9, value=total)
        total_cell.font = Font(bold=True, name="Calibri")
        total_cell.number_format = "#,##0.00"

        # Column widths
        widths = [18, 30, 40, 18, 18, 8, 15, 10, 18]
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[chr(64 + i)].width = w

        filename = f"bom_{metadata.get('rfp_id', 'draft')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        output_path = self.output_dir / filename
        wb.save(str(output_path))
        logger.info("Created BOM spreadsheet: %s", output_path)
        return str(output_path)
