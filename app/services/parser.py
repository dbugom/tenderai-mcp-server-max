"""Parser service — PDF and DOCX ingestion for RFP documents."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class ParserService:
    """Extracts text and tables from PDF and DOCX files."""

    def __init__(self, data_dir: Path):
        self.data_dir = data_dir

    async def parse_file(self, file_path: str) -> dict:
        """Auto-detect format and parse. Returns {text, tables, page_count, format}."""
        p = Path(file_path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = p.suffix.lower()
        if ext == ".pdf":
            return await self.parse_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return await self.parse_docx(file_path)
        elif ext in (".xlsx", ".xls"):
            return await self.parse_xlsx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    async def parse_pdf(self, file_path: str) -> dict:
        """Extract text and tables from a PDF using pdfplumber."""
        import pdfplumber

        text_parts: list[str] = []
        tables: list[list] = []

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)

        logger.info("Parsed PDF: %s (%d pages, %d tables)", file_path, page_count, len(tables))
        return {
            "text": "\n\n".join(text_parts),
            "tables": tables,
            "page_count": page_count,
            "format": "pdf",
        }

    async def parse_docx(self, file_path: str) -> dict:
        """Extract text and tables from a DOCX using python-docx."""
        from docx import Document

        doc = Document(file_path)
        text_parts: list[str] = []
        tables: list[list] = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            tables.append(table_data)

        logger.info("Parsed DOCX: %s (%d paragraphs, %d tables)", file_path, len(text_parts), len(tables))
        return {
            "text": "\n\n".join(text_parts),
            "tables": tables,
            "page_count": None,
            "format": "docx",
        }

    async def parse_xlsx(self, file_path: str) -> dict:
        """Extract data from an XLSX file using openpyxl."""
        from openpyxl import load_workbook

        wb = load_workbook(file_path, data_only=True)
        text_parts: list[str] = []
        tables: list[list] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"Sheet: {sheet_name}")
            sheet_data = []
            for row in ws.iter_rows(values_only=True):
                str_row = [str(cell) if cell is not None else "" for cell in row]
                sheet_data.append(str_row)
                text_parts.append(" | ".join(str_row))
            tables.append(sheet_data)

        logger.info("Parsed XLSX: %s (%d sheets)", file_path, len(wb.sheetnames))
        return {
            "text": "\n".join(text_parts),
            "tables": tables,
            "page_count": None,
            "format": "xlsx",
        }
